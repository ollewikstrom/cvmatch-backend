import re
from datetime import datetime
from docx import Document
import json

def detect_applicant_name(document):
    """
    Extracts the applicant's first and last name from Table 1, Row 2.
    """
    try:
        table = document.tables[0]  # Table 1
        name_row = table.rows[1]  # Row 2
        name = name_row.cells[1].text.strip()  # Extract the middle cell
        first_name, last_name = name.split(" ", 1)
        return first_name, last_name
    except (IndexError, ValueError):
        return None, None  # Return None if name extraction fails


def extract_introduction(document):
    """
    Extracts the introduction from Table 2, Row 1.
    """
    try:
        table = document.tables[1]  # Table 2
        introduction_row = table.rows[0]  # Row 1
        introduction = introduction_row.cells[0].text.strip()  # Extract the first cell
        return introduction
    except IndexError:
        return None  # Return None if introduction extraction fails


def extract_education(document):
    """
    Extracts the education section from Table 3, Row 3.
    """
    try:
        table = document.tables[2]  # Table 3
        education_row = table.rows[2]  # Row 3
        education = education_row.cells[0].text.strip()  # Extract the first cell
        return education
    except IndexError:
        return None  # Return None if education extraction fails


def detect_company_names(document):
    company_names = set()
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            match = re.search(r"(Kund|Client):\s*(.+)", row_text)
            if match:
                company_names.add(match.group(2).strip())

    for para in document.paragraphs:
        match = re.search(r"(Kund|Client):\s*(.+)", para.text)
        if match:
            company_name = match.group(2).strip()
            if company_name:
                company_names.add(company_name)
    return company_names


def replace_prohibited_terms(json_data, prohibited_terms):
    """Replace prohibited terms in specific fields with '[REDACTED]'."""
    # Filter out empty or invalid prohibited terms
    valid_terms = [term for term in prohibited_terms if term.strip()]
    if not valid_terms:  # If no valid terms, skip replacement
        return json_data

    # Build regex pattern to match whole words only
    pattern = r"\b(" + "|".join(re.escape(term) for term in valid_terms) + r")\b"

    fields_to_check = ["description", "applied_skills", "customer"]  # Fields to clean
    for key in fields_to_check:
        value = json_data.get(key, "")
        
        if isinstance(value, str):  # For string fields
            value = re.sub(pattern, "[REDACTED]", value)
            json_data[key] = value

        elif isinstance(value, list):  # For list fields (e.g., skills)
            json_data[key] = [
                re.sub(pattern, "[REDACTED]", item) for item in value
            ]

    return json_data




def convert_to_json(input_text, prohibited_terms):
    """Convert a section of text into structured JSON data."""
    # Normalize input text - preserve line breaks for multi-line matches
    input_text = "\n".join(input_text.splitlines())

    # Match 'Roll:' or 'Role:' followed by optional pipes and spaces
    role_match = re.search(r"(Roll|Role):\s*\|?\s*(.+)", input_text)
    role = role_match.group(2).strip() if role_match else ""

    # Match 'Kund:' or 'Client:' followed by optional pipes and spaces
    customer_match = re.search(r"(Kund|Client):\s*\|?\s*(.+)", input_text)
    customer = customer_match.group(2).strip() if customer_match else ""

    # Match 'Period:' for date extraction
    period_match = re.search(r"Period:\s*\|?\s*([\d\-–\s]+)", input_text)
    if period_match:
        date_range = period_match.group(1).split("–")
        date_start = (
            datetime.strptime(date_range[0].strip(), "%Y-%m").strftime("%m/%Y")
            if date_range[0].strip()
            else ""
        )
        date_end = (
            datetime.strptime(date_range[1].strip(), "%Y-%m").strftime("%m/%Y")
            if len(date_range) > 1 and date_range[1].strip()
            else ""
        )
    else:
        date_start = ""
        date_end = ""

    # Extract description
    description_match = re.search(r"(Beskrivning|Description)\s*(.+?)(?=\nExpertis|Expertise|\Z)", input_text, re.DOTALL)
    description = description_match.group(2).strip() if description_match else ""

    # Extract expertise/skills
    expertise_match = re.search(r"(Expertis|Expertise)\s*\|(.+)", input_text, re.DOTALL)
    expertise = parse_skills(expertise_match.group(2)) if expertise_match else []

    output_json = {
        "role": role,
        "customer": customer,
        "date_start": date_start,
        "date_end": date_end,
        "description": description,
        "expertise": expertise,
    }

    # Replace prohibited terms with '[REDACTED]'
    output_json = replace_prohibited_terms(output_json, prohibited_terms)

    return output_json


def parse_skills(skill_string):
    """Parse the skills line into a list of skills."""
    return [skill.strip() for skill in skill_string.split('|') if skill.strip()]


def find_line_with_pipe(input_string):
    """Find the line containing '|' which lists the skills."""
    for line in input_string.splitlines():
        if "|" in line:
            return line.strip()
    return ""

def extract_section_by_heading_and_tables(document, heading):
    """Extracts content under a specific heading from both paragraphs and tables."""
    content = []
    capture = False

    # Extract paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if text.startswith(heading):
            capture = True
        elif capture and text:  # Collect text until the next heading or empty line
            if text.isupper():  # Assuming headings are uppercase
                break
            content.append(text)

    # Extract from tables
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            if heading.lower() in row_text.lower():
                capture = True
                continue
            if capture:
                if row_text.lower().startswith("summary") or row_text.lower().startswith("education"):
                    break  # Stop capturing if another section starts
                content.append(row_text)

    return "\n".join(content)

def extract_sections_from_tables(document):
    """Extract sections from tables that contain rows starting with 'Roll'."""
    sections = []
    current_section = []
    seen_rows = set()  # Track unique row content
    capture = False

    for table in document.tables:
        for row in table.rows:            # Extract unique content from cells
            cells = list(dict.fromkeys([cell.text.strip() for cell in row.cells if cell.text.strip()]))
            row_text = " | ".join(cells).strip()


            # Skip empty rows or rows already processed
            if not row_text or row_text in seen_rows:
                continue
            seen_rows.add(row_text)

            # Check for the start of a new section
            if re.match(r"^(Roll|Role):", row_text):
               
                if current_section:  # Save the previous section
                    sections.append("\n".join(current_section))
                current_section = [row_text]
                capture = True
            elif capture:
                current_section.append(row_text)

        # End of table - save the last section
        if current_section:
            sections.append("\n".join(current_section))
            current_section = []
            seen_rows.clear()  # Reset for the next table

    return sections


def parse_docx_to_json(cv_file, filename):
    """Read DOCX and convert to filtered JSON."""
    document = Document(cv_file)

    # Extract summary under "Introduktion"
    summary = extract_introduction(document)

    # Extract education under "Utbildning"
    education = extract_education(document)

    # Step 1: Detect applicant's name and company names
    first_name, last_name = detect_applicant_name(document)
    company_names = detect_company_names(document)
    prohibited_terms = {first_name, last_name}.union(company_names)

    # Step 2: Extract sections starting with 'Roll:' from tables
    sections = extract_sections_from_tables(document)

    # Step 3: Process each section and filter results
    all_assignments = []
    for section in sections:
        parsed_json = convert_to_json(section, prohibited_terms)
        if parsed_json:  # Only include non-filtered sections
            all_assignments.append(parsed_json)
    output_json = {"title": filename, "introduction": summary, "education": education, "assignments": all_assignments}

    return output_json


